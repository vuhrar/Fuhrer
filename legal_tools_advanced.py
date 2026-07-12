# legal_tools_advanced.py
"""
الأدوات القانونية المتقدمة — النسخة الاحترافية v2.0
إصلاح كامل لجميع الأعطال + ميزات جديدة
المبادئ: Single Responsibility, Clean Contracts, Full Error Handling
"""

import re
import json
from typing import List, Dict, Optional, Callable, Tuple
from datetime import datetime, timedelta  # ← إصلاح: استيراد timedelta
from io import BytesIO

from legal_database import get_legal_database

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================================
# 1. نظام تصنيف آلي للقضايا
# ============================================================
class LegalCaseClassifier:
    """تصنيف آلي لقضايا القانون العمالي باستخدام كلمات رئيسية."""

    CATEGORIES = {
        "نهاية الخدمة": ["نهاية خدمة", "مكافأة نهاية خدمة", "مادة 84", "علاوة نهاية"],
        "فصل تعسفي": ["فصل تعسفي", "فصل بدون سبب", "إنهاء تعسفي", "مادة 77"],
        "الأجور والرواتب": ["راتب", "أجر", "زيادة", "خصم", "مستحقات", "مادة 90", "تأخير راتب"],
        "عقد العمل": ["عقد", "عقد عمل", "تجديد", "مدة", "مادة 78", "إشعار"],
        "ساعات العمل": ["ساعات", "عمل إضافي", "فترة راحة", "مادة 98", "مادة 106"],
        "التأمينات الاجتماعية": ["تأمينات", "GOSI", "معاش", "مادة 130"],
        "المنازعات العمالية": ["منازعة", "دعوى", "محكمة", "مادة 209"],
        "السلامة المهنية": ["سلامة", "مخاطر", "حماية", "مادة 121"],
        "الإجازات": ["إجازة", "إجازة سنوية", "مادة 109", "مادة 113"],
        "التأديب": ["تأديب", "إنذار", "عقوبة", "مادة 155"],
        "إصابات العمل": ["إصابة عمل", "حادث عمل", "مادة 131"],
        "السعودة": ["سعودة", "نطاقات", "توطين"],
    }

    def classify(self, text: str) -> Tuple[str, float, List[str]]:
        """
        تصنيف نص قضية إلى فئة قانونية.
        
        Returns:
            Tuple[category, confidence, matched_keywords]
        """
        text_lower = text.lower()
        scores: Dict[str, float] = {}
        matched: Dict[str, List[str]] = {}

        for category, keywords in self.CATEGORIES.items():
            found = [kw for kw in keywords if kw.lower() in text_lower]
            if found:
                scores[category] = len(found) / len(keywords)
                matched[category] = found

        if not scores:
            return "عام", 0.0, []

        best = max(scores.items(), key=lambda x: x[1])
        return best[0], round(best[1], 2), matched.get(best[0], [])

    def classify_with_ai(self, text: str, call_ai_fn: Callable) -> str:
        """تصنيف باستخدام AI لزيادة الدقة."""
        categories_str = "، ".join(self.CATEGORIES.keys())
        prompt = (
            f"صنّف النص القانوني التالي إلى إحدى هذه الفئات فقط:\n"
            f"{categories_str}\n\n"
            f"النص:\n{text[:500]}\n\n"
            f"أجب بكلمة واحدة فقط: اسم الفئة."
        )
        try:
            result = call_ai_fn(prompt, [], "أنت خبير في تصنيف القضايا القانونية السعودية.").strip()
            for cat in self.CATEGORIES:
                if cat in result:
                    return cat
            return "عام"
        except Exception:
            return "عام"


# ============================================================
# 2. محرك البحث القانوني
# ============================================================
class LegalSearchEngine:
    """محرك بحث متقدم في نصوص القانون العمالي السعودي."""

    def __init__(self):
        self.db = get_legal_database()

    def search(self, query: str, category: Optional[str] = None, max_results: int = 10) -> List[Dict]:
        """بحث في قاعدة البيانات القانونية."""
        results = self.db.search_laws(query, category)
        query_lower = query.lower()
        scored = []
        for law in results:
            text = law.get("text", "").lower()
            title = law.get("title", "").lower()
            keywords = " ".join(law.get("keywords", [])).lower()
            score = (
                text.count(query_lower) * 2 +
                title.count(query_lower) * 5 +
                keywords.count(query_lower) * 3
            )
            scored.append({**law, "score": min(1.0, score / 10)})
        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[:max_results]

    def search_with_ai(self, query: str, call_ai_fn: Callable, max_results: int = 5) -> List[Dict]:
        """بحث مُعزَّز بالذكاء الاصطناعي."""
        # أولاً: بحث محلي
        local = self.search(query, max_results=max_results)
        if local:
            return local
        # ثانياً: سؤال الذكاء الاصطناعي
        prompt = (
            f"ابحث في نظام العمل السعودي عن: {query}\n"
            f"أعطني المواد القانونية ذات الصلة مع نصوصها."
        )
        try:
            response = call_ai_fn(prompt, [], "أنت خبير في القانون العمالي السعودي.")
            return [{"title": "نتيجة AI", "text": response, "article": "AI", "category": "عام", "score": 1.0}]
        except Exception:
            return []


# ============================================================
# 3. أداة التفنيد القانوني
# ============================================================
class LegalRebuttalGenerator:
    """توليد تفنيد قانوني احترافي للادعاءات."""

    def generate(self, claim: str, case_details: str, call_ai_fn: Callable) -> Dict:
        """
        توليد تفنيد قانوني للادعاء.
        
        Returns:
            Dict with keys: analysis, legal_articles, arguments, conclusion, error
        """
        prompt = (
            f"أنت محامٍ سعودي خبير بالقانون العمالي بخبرة 20 عاماً.\n\n"
            f"الادعاء:\n{claim}\n\n"
            f"تفاصيل القضية:\n{case_details}\n\n"
            f"قم بتوليد تفنيد قانوني مفصل يشمل:\n"
            f"1. تحليل الادعاء ونقاط ضعفه\n"
            f"2. المواد القانونية الداعمة للدفاع\n"
            f"3. الحجج الدفاعية القوية\n"
            f"4. الاستنتاج القانوني\n\n"
            f"استند إلى نصوص نظام العمل السعودي بدقة."
        )
        try:
            response = call_ai_fn(prompt, [], "أنت محامٍ سعودي متخصص في القانون العمالي.")
            return {
                "analysis": response,
                "legal_articles": [],
                "arguments": [],
                "conclusion": "",
                "error": None
            }
        except Exception as e:
            return {
                "analysis": "",
                "legal_articles": [],
                "arguments": [],
                "conclusion": "",
                "error": str(e)
            }


# ============================================================
# 4. توليد المستندات القانونية
# ============================================================
class LegalDocumentGenerator:
    """توليد مستندات قانونية احترافية بصيغة DOCX."""

    def generate_lawsuit(self, case_data: Dict) -> Optional[bytes]:
        """توليد دعوى عمالية بصيغة DOCX."""
        if not DOCX_AVAILABLE:
            return None
        try:
            doc = Document()
            # إعداد الخط
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            # العنوان
            title = doc.add_heading("دعوى عمالية", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"إلى: لجنة تسوية المنازعات العمالية في {case_data.get('court_city', '...')}")
            doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"الرقم المرجعي: {case_data.get('case_number', 'جديد')}")
            doc.add_paragraph("")

            # بيانات المدعي
            doc.add_heading("بيانات المدعي:", level=2)
            doc.add_paragraph(f"الاسم: {case_data.get('plaintiff_name', '...')}")
            doc.add_paragraph(f"الجنسية: {case_data.get('plaintiff_nationality', '...')}")
            doc.add_paragraph(f"رقم الهوية: {case_data.get('plaintiff_id', '...')}")

            # بيانات المدعى عليه
            doc.add_heading("بيانات المدعى عليه:", level=2)
            doc.add_paragraph(f"الاسم: {case_data.get('defendant_name', '...')}")
            doc.add_paragraph(f"العنوان: {case_data.get('defendant_address', '...')}")

            # موضوع الدعوى
            doc.add_heading("موضوع الدعوى:", level=2)
            doc.add_paragraph(case_data.get('case_subject', ''))

            # الوقائع
            doc.add_heading("وقائع الدعوى:", level=2)
            for i, fact in enumerate(case_data.get('facts', []), 1):
                doc.add_paragraph(f"{i}. {fact}")

            # المطالبات
            doc.add_heading("المطالبات:", level=2)
            for i, claim in enumerate(case_data.get('claims', []), 1):
                doc.add_paragraph(f"{i}. {claim}")

            # التوقيع
            doc.add_paragraph("")
            doc.add_paragraph("المدعي / التوقيع: ___________________")
            doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception:
            return None

    def generate_defense(self, case_data: Dict) -> Optional[bytes]:
        """توليد مذكرة دفاع قانونية."""
        if not DOCX_AVAILABLE:
            return None
        try:
            doc = Document()
            doc.add_heading("مذكرة دفاع قانونية", level=1)
            doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_heading("بيانات المدعى عليه:", level=2)
            doc.add_paragraph(f"الاسم: {case_data.get('defendant_name', '...')}")
            doc.add_heading("موضوع الدفاع:", level=2)
            doc.add_paragraph(case_data.get('defense_subject', ''))
            doc.add_heading("حجج الدفاع:", level=2)
            for i, arg in enumerate(case_data.get('arguments', []), 1):
                doc.add_paragraph(f"{i}. {arg}")
            doc.add_heading("الطلبات:", level=2)
            for i, req in enumerate(case_data.get('requests', []), 1):
                doc.add_paragraph(f"{i}. {req}")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception:
            return None


# ============================================================
# 5. حاسبة الاستحقاقات — إصلاح كامل للعقد
# ============================================================
class LegalEntitlementsCalculator:
    """حاسبة متقدمة لاستحقاقات القانون العمالي السعودي."""

    def calculate_eosb(
        self,
        basic_salary: float,
        total_salary: float,
        years_of_service: float,
        is_arbitrary: bool = True,
        delay_months: int = 0,
        is_saudi: bool = True,
        resignation: bool = False
    ) -> Dict:
        """
        حساب مكافأة نهاية الخدمة وفقاً للمادة 84.
        
        Returns:
            Dict with keys: basic_salary, total_salary, years_of_service,
                           totals (grand_total, eosb_total, arbitrary_total, delay_total),
                           details (dict of components),
                           summary (list of strings)
        """
        # حساب المكافأة الأساسية
        if resignation:
            # نسب الاستقالة حسب المادة 84
            if years_of_service < 2:
                eosb = 0.0
                eosb_note = "لا تستحق مكافأة (أقل من سنتين)"
            elif years_of_service < 5:
                eosb = (basic_salary / 2) * years_of_service * (1/3)
                eosb_note = "ثلث المكافأة (استقالة 2-5 سنوات)"
            elif years_of_service < 10:
                eosb = (basic_salary / 2) * min(years_of_service, 5) + basic_salary * max(0, years_of_service - 5)
                eosb = eosb * 0.5
                eosb_note = "نصف المكافأة (استقالة 5-10 سنوات)"
            else:
                y5 = min(years_of_service, 5)
                y_plus = max(0, years_of_service - 5)
                eosb = (basic_salary / 2) * y5 + basic_salary * y_plus
                eosb_note = "المكافأة كاملة (استقالة بعد 10 سنوات)"
        else:
            # فصل أو انتهاء عقد: مكافأة كاملة
            y5 = min(years_of_service, 5)
            y_plus = max(0, years_of_service - 5)
            eosb = (basic_salary / 2) * y5 + basic_salary * y_plus
            eosb_note = "المكافأة كاملة (فصل أو انتهاء عقد)"

        # تعويض الفصل التعسفي (مادة 77)
        arb = 0.0
        if is_arbitrary and not resignation:
            arb_months = min(12, max(3, int(years_of_service)))
            arb = total_salary * arb_months
        arb_note = f"تعويض {min(12, max(3, int(years_of_service)))} شهر" if arb > 0 else "لا ينطبق"

        # تعويض تأخير الراتب (مادة 90)
        delay = total_salary * 0.05 * delay_months if delay_months > 0 else 0.0
        delay_note = f"5% × {delay_months} شهر تأخير" if delay > 0 else "لا ينطبق"

        grand_total = eosb + arb + delay

        # ← إصلاح: إضافة مفتاح details الذي تتوقعه الواجهة
        details = {
            "eosb": {
                "description": f"مكافأة نهاية الخدمة ({eosb_note})",
                "amount": round(eosb, 2),
                "formula": f"أجر أساسي {basic_salary:,.0f} × {years_of_service:.1f} سنة"
            },
            "arbitrary": {
                "description": f"تعويض الفصل التعسفي ({arb_note})",
                "amount": round(arb, 2),
                "formula": f"أجر إجمالي {total_salary:,.0f} × أشهر التعويض"
            },
            "delay": {
                "description": f"تعويض تأخير الراتب ({delay_note})",
                "amount": round(delay, 2),
                "formula": f"5% × {total_salary:,.0f} × {delay_months} شهر"
            },
        }

        return {
            "basic_salary": round(basic_salary, 2),
            "total_salary": round(total_salary, 2),
            "years_of_service": round(years_of_service, 2),
            "is_saudi": is_saudi,
            "resignation": resignation,
            "totals": {
                "eosb_total": round(eosb, 2),
                "arbitrary_total": round(arb, 2),
                "delay_total": round(delay, 2),
                "grand_total": round(grand_total, 2),
            },
            "details": details,  # ← مُصلَح
            "summary": [
                f"مكافأة نهاية الخدمة: {eosb:,.2f} ريال",
                f"تعويض الفصل التعسفي: {arb:,.2f} ريال" if arb > 0 else None,
                f"تعويض تأخير الراتب: {delay:,.2f} ريال" if delay > 0 else None,
                f"**الإجمالي: {grand_total:,.2f} ريال**",
            ]
        }

    def calculate_vacation_balance(
        self,
        start_date: str,
        end_date: str = "",
        annual_vacation_days: int = 21
    ) -> Dict:
        """حساب رصيد الإجازة السنوية."""
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
            end = datetime.strptime(end_date, "%Y-%m-%d") if end_date else datetime.now()
            delta = end - start
            years = delta.days / 365.25
            full_years = int(years)
            remaining_fraction = years - full_years

            # بعد 5 سنوات: 30 يوم بدلاً من 21
            if years >= 5:
                annual_vacation_days = 30

            earned = years * annual_vacation_days
            full_years_vacation = full_years * annual_vacation_days
            remaining_vacation = remaining_fraction * annual_vacation_days

            return {
                "start_date": start_date,
                "end_date": end_date or datetime.now().strftime("%Y-%m-%d"),
                "total_years": round(years, 2),
                "annual_vacation_days": annual_vacation_days,
                "earned_vacation": round(earned, 1),
                "full_years_vacation": round(full_years_vacation, 1),
                "remaining_vacation": round(remaining_vacation, 1),
                "cash_value_note": "يُحسب بقسمة الراتب الشهري على 30 × عدد الأيام",
            }
        except Exception as e:
            return {"error": str(e), "earned_vacation": 0, "full_years_vacation": 0, "remaining_vacation": 0}

    def calculate_overtime(
        self,
        hourly_rate: float,
        overtime_hours: float,
        overtime_rate: float = 1.5
    ) -> Dict:
        """حساب أجر العمل الإضافي (مادة 106)."""
        overtime_pay = hourly_rate * overtime_hours * overtime_rate
        return {
            "hourly_rate": round(hourly_rate, 2),
            "overtime_hours": overtime_hours,
            "overtime_rate": overtime_rate,
            "overtime_pay": round(overtime_pay, 2),
            "legal_basis": "المادة 106 من نظام العمل — 150% من أجر الساعة العادية",
        }


# ============================================================
# 6. نظام تتبع المواعيد — إصلاح timedelta
# ============================================================
class LegalDeadlinesTracker:
    """نظام تتبع المواعيد القانونية الهامة."""

    DEADLINES = {
        "نهاية الخدمة": {
            "description": "مهلة رفع دعوى نهاية الخدمة",
            "days": 365, "article": "222",
            "action": "رفع دعوى إلى المحكمة العمالية خلال سنة من انتهاء العقد"
        },
        "فصل تعسفي": {
            "description": "مهلة رفع دعوى ضد الفصل التعسفي",
            "days": 365, "article": "77",
            "action": "رفع دعوى إلى المحكمة العمالية خلال سنة"
        },
        "تأخير راتب": {
            "description": "مهلة امتناع العامل عن العمل بسبب تأخير الراتب",
            "days": 7, "article": "90",
            "action": "يحق للعامل الامتناع عن العمل بعد 7 أيام من تأخير الراتب"
        },
        "إشعار إنهاء العقد": {
            "description": "مهلة إشعار صاحب العمل قبل إنهاء العقد",
            "days": 60, "article": "78",
            "action": "تقديم إشعار كتابي قبل 60 يوماً (للعمال الشهريين)"
        },
        "إصابة عمل": {
            "description": "مهلة الإبلاغ عن إصابة عمل",
            "days": 7, "article": "131",
            "action": "إبلاغ صاحب العمل ووزارة الموارد البشرية خلال 7 أيام"
        },
        "الطعن في القرار التأديبي": {
            "description": "مهلة الطعن في قرار تأديبي",
            "days": 30, "article": "155",
            "action": "تقديم اعتراض كتابي لصاحب العمل خلال 30 يوماً"
        },
    }

    def check_deadline(self, case_type: str, case_date: str) -> Dict:
        """
        التحقق من الموعد القانوني لقضية معينة.
        
        Returns:
            Dict with keys: case_type, description, article, case_date,
                           due_date, days_remaining, status, color, message, action
        """
        if case_type not in self.DEADLINES:
            return {
                "error": f"نوع القضية '{case_type}' غير معروف",
                "message": "نوع قضية غير معروف",
                "status": "unknown", "color": "#666666",
                "days_remaining": 0, "due_date": "", "case_date": case_date,
                "article": "", "action": ""
            }

        deadline_info = self.DEADLINES[case_type]
        try:
            case_date_obj = datetime.strptime(case_date, "%Y-%m-%d")
        except ValueError:
            case_date_obj = datetime.now()

        # ← إصلاح: timedelta مستورد الآن
        due_date = case_date_obj + timedelta(days=deadline_info["days"])
        today = datetime.now()
        days_remaining = (due_date - today).days

        if days_remaining < 0:
            status, color = "منتهية", "#ff5a5a"
            message = f"انتهت المهلة منذ {-days_remaining} يوم"
        elif days_remaining <= 7:
            status, color = "عاجلة", "#ff8c00"
            message = f"مهلة عاجلة جداً — {days_remaining} يوم متبقٍ فقط!"
        elif days_remaining <= 30:
            status, color = "قريبة", "#ffb400"
            message = f"المهلة قريبة — {days_remaining} يوم متبقٍ"
        else:
            status, color = "متاحة", "#4ade80"
            message = f"لديك وقت كافٍ — {days_remaining} يوم متبقٍ"

        return {
            "case_type": case_type,
            "description": deadline_info["description"],
            "article": deadline_info["article"],
            "case_date": case_date,
            "due_date": due_date.strftime("%Y-%m-%d"),
            "days_remaining": days_remaining,
            "status": status,
            "color": color,
            "message": message,
            "action": deadline_info["action"],
        }

    @property
    def deadlines(self) -> Dict:
        """خاصية للوصول إلى قاموس المواعيد."""
        return self.DEADLINES


# ============================================================
# 7. نظام كشف الزلات القانونية — إصلاح مفاتيح الإخراج
# ============================================================
class LegalSlipDetector:
    """نظام كشف الزلات القانونية في المراسلات."""

    PATTERNS = [
        (r"(أكرهنا|أجبرنا|إكراه|ضغط عليه|تهديد)", "إكراه على الاستقالة", "danger", "77",
         "الإكراه على الاستقالة يُعدّ فصلاً تعسفياً — احتفظ بهذا الدليل"),
        (r"(لا نعترف|لم نقر|لا نقر)\s*(بـ|ب)?\s*(حقوق|مستحقات|علاوة)", "رفض الاعتراف بحقوق العامل", "danger", "84",
         "رفض المستحقات القانونية مخالفة صريحة للمادة 84 — وثّق هذا الرفض"),
        # نمط أوسع: فصل / إنهاء بدون سبب أو تحقيق أو إشعار
        (r"(فصل|إنهاء|إخراج|\u0641\u0635\u0644|\u0625\u0646\u0647\u0627\u0621)[^\n]{0,30}(بدون|دون|\u0628\u062f\u0648\u0646|\u062f\u0648\u0646)[^\n]{0,20}(سبب|تحقيق|إشعار|\u0633\u0628\u0628|\u062a\u062d\u0642\u064a\u0642|\u0625\u0634\u0639\u0627\u0631)", "فصل بدون إجراءات قانونية", "danger", "77",
         "الفصل بدون تحقيق مخالفة للمادة 155 — يستوجب التعويض"),
        # نمط إضافي: فصل فوري / فصل مباشر
        (r"(فصل|إنهاء)[^\n]{0,15}(فوراً|فوري|مباشرة|فور)", "فصل فوري بدون إشعار", "danger", "75",
         "الفصل الفوري بدون سبب مشروع يستوجب تعويضاً وفق المادة 77"),
        (r"(خطأ|خطئاً|غلطة)\s*(من|في)\s*(الجهة|صاحب العمل|الشركة|الإدارة)", "إقرار بخطأ من صاحب العمل", "danger", "77",
         "هذا الإقرار دليل قوي لصالحك — احتفظ به"),
        (r"(حرمان|محروم)\s*(من|عن)\s*(الأجر|الراتب|المستحقات)", "حرمان من مستحقات مالية", "danger", "90",
         "الحرمان من الأجر مخالفة للمادة 90 — يحق لك المطالبة بالتعويض"),
        (r"(لا يستحق|لم يستحق|غير مستحق)\s*(علاوة|مكافأة|تعويض)", "رفض استحقاق قانوني", "warn", "84",
         "تحقق من حساباتك — قد يكون هذا الرفض غير قانوني"),
        (r"(تأخير|تأخر)\s*(في|ب)?\s*(دفع|صرف)\s*(الراتب|الأجر)", "تأخير في دفع الأجر", "warn", "90",
         "تأخير الراتب أكثر من 7 أيام يُخوّلك الامتناع عن العمل مع الاحتفاظ بحقوقك"),
        (r"(إنذار|إنذارات)\s*(بدون|دون|من غير)\s*(سبب|مبرر|تحقيق)", "إنذارات بدون مبرر", "warn", "155",
         "الإنذار بدون تحقيق مخالف للمادة 155 — يمكن الطعن فيه"),
        (r"(تمييز|تفريق|معاملة مختلفة)\s*(بين|في)\s*(العمال|الموظفين)", "تمييز بين العمال", "warn", "3",
         "التمييز بين العمال مخالف لمبادئ نظام العمل"),
    ]

    def detect(self, text: str) -> List[Dict]:
        """
        كشف الزلات في النص.
        
        Returns:
            List of Dicts with keys: msg, level, snippet, article, suggestion
            (← إصلاح: تغيير 'message' إلى 'msg' وإضافة 'suggestion')
        """
        findings = []
        seen = set()
        for pattern, msg, level, article, suggestion in self.PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches and msg not in seen:
                seen.add(msg)
                snippet = matches[0] if isinstance(matches[0], str) else " ".join(matches[0])
                findings.append({
                    "msg": msg,                    # ← مُصلَح (كان 'message')
                    "level": level,
                    "snippet": snippet[:120].strip(),
                    "article": article,
                    "suggestion": suggestion,      # ← مُضاف جديد
                })
        return findings


# ============================================================
# 8. نظام تحليل قوة القضية — إصلاح كامل للعقد
# ============================================================
class CaseStrengthAnalyzer:
    """نظام تحليل قوة القضية القانونية."""

    def quick_analyze(self, case_text: str, category: str = "عام") -> Dict:
        """
        تحليل سريع بدون AI.
        
        Args:
            case_text: نص تفاصيل القضية (← إصلاح: تغيير التوقيع)
            category: تصنيف القضية
        
        Returns:
            Dict with keys: score, strength, success_probability, analysis, recommendations
        """
        # تقييم بناءً على التصنيف
        category_scores = {
            "نهاية الخدمة": 7, "فصل تعسفي": 8, "الأجور والرواتب": 7,
            "إصابات العمل": 8, "المنازعات العمالية": 6, "التأديب": 6,
            "عقد العمل": 5, "ساعات العمل": 6, "الإجازات": 6,
        }
        base_score = category_scores.get(category, 5)

        # تحليل نص القضية لتحديد قوة الأدلة
        evidence_keywords = ["وثيقة", "عقد", "إشعار", "شاهد", "بريد", "رسالة", "تسجيل", "صورة", "دليل"]
        evidence_count = sum(1 for kw in evidence_keywords if kw in case_text)
        evidence_bonus = min(2, evidence_count * 0.5)

        # تحليل الكلمات القانونية
        legal_keywords = ["مادة", "نظام العمل", "حق", "مستحق", "قانون", "محكمة"]
        legal_count = sum(1 for kw in legal_keywords if kw in case_text)
        legal_bonus = min(1, legal_count * 0.3)

        score = min(10, max(1, round(base_score + evidence_bonus + legal_bonus)))

        if score >= 8:
            strength = "قوية جداً"
            probability = f"{score * 10}%"
        elif score >= 6:
            strength = "قوية"
            probability = f"{score * 9}%"
        elif score >= 4:
            strength = "متوسطة"
            probability = f"{score * 8}%"
        else:
            strength = "ضعيفة"
            probability = f"{score * 7}%"

        # توليد التحليل والتوصيات
        analysis_text = (
            f"بناءً على تصنيف القضية كـ '{category}'، "
            f"تبدو القضية {strength} مع احتمال نجاح يُقدَّر بـ {probability}. "
            f"تم رصد {evidence_count} مؤشر على وجود أدلة داعمة."
        )

        recommendations = [
            "اجمع جميع المستندات المتعلقة بالقضية (عقد العمل، الرواتب، المراسلات)",
            "توثيق جميع الوقائع بالتواريخ والأدلة الكتابية",
            "استشر محامياً متخصصاً في القانون العمالي قبل رفع الدعوى",
        ]

        if category == "فصل تعسفي":
            recommendations.append("تأكد من استيفاء شروط التعويض وفقاً للمادة 77")
        elif category == "نهاية الخدمة":
            recommendations.append("احسب مستحقاتك بدقة وفقاً للمادة 84 قبل التفاوض")
        elif category == "الأجور والرواتب":
            recommendations.append("احتفظ بكشوف الراتب وسجلات الدفع كأدلة")

        return {
            "score": score,
            "strength": strength,
            "success_probability": probability,
            "analysis": analysis_text,          # ← مُضاف جديد
            "recommendations": recommendations,  # ← مُضاف جديد
        }

    def analyze_with_ai(self, case_text: str, call_ai_fn: Callable) -> Dict:
        """تحليل متعمق بالذكاء الاصطناعي."""
        prompt = (
            f"أنت خبير قانوني في نظام العمل السعودي.\n\n"
            f"حلّل قوة هذه القضية:\n{case_text[:1000]}\n\n"
            f"قدّم:\n"
            f"1. درجة قوة القضية من 1 إلى 10\n"
            f"2. احتمال النجاح كنسبة مئوية\n"
            f"3. نقاط القوة الرئيسية\n"
            f"4. نقاط الضعف المحتملة\n"
            f"5. التوصيات العملية\n"
            f"6. المواد القانونية ذات الصلة"
        )
        try:
            response = call_ai_fn(prompt, [], "أنت خبير قانوني في نظام العمل السعودي.")
            # استخراج الدرجة من الرد
            score_match = re.search(r'(\d+)\s*(?:/|من)\s*10', response)
            score = int(score_match.group(1)) if score_match else 6
            score = min(10, max(1, score))
            return {
                "score": score,
                "strength": "قوية" if score >= 6 else "متوسطة",
                "success_probability": f"{score * 10}%",
                "analysis": response,
                "recommendations": [],
            }
        except Exception as e:
            return {"score": 5, "strength": "متوسطة", "success_probability": "50%",
                    "analysis": f"خطأ في التحليل: {str(e)}", "recommendations": []}


# ============================================================
# 9. استخراج المعلومات
# ============================================================
class LegalInfoExtractor:
    """استخراج معلومات قانونية من النصوص."""

    def extract_dates(self, text: str) -> List[str]:
        patterns = [r'\d{1,2}/\d{1,2}/\d{2,4}', r'\d{4}-\d{2}-\d{2}', r'\d{1,2}-\d{1,2}-\d{2,4}']
        dates = []
        for p in patterns:
            dates.extend(re.findall(p, text))
        return list(set(dates))[:10]

    def extract_amounts(self, text: str) -> List[str]:
        patterns = [
            r'\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*(?:ريال|SR|ر\.س)',
            r'\d{1,3}(?:,\d{3})+(?:\.\d+)?'
        ]
        amounts = []
        for p in patterns:
            amounts.extend(re.findall(p, text))
        return list(set(amounts))[:10]

    def extract_articles(self, text: str) -> List[str]:
        patterns = [r'المادة\s+\d+', r'م\s*\d+', r'مادة\s+\d+']
        articles = []
        for p in patterns:
            articles.extend(re.findall(p, text, re.IGNORECASE))
        return list(set(articles))[:10]

    def extract_all(self, text: str) -> Dict:
        return {
            "dates": self.extract_dates(text),
            "amounts": self.extract_amounts(text),
            "articles": self.extract_articles(text),
        }


# ============================================================
# Singleton Instances — نسخ عالمية جاهزة للاستخدام
# ============================================================
legal_classifier    = LegalCaseClassifier()
legal_search        = LegalSearchEngine()
rebuttal_generator  = LegalRebuttalGenerator()
document_generator  = LegalDocumentGenerator()
entitlements_calculator = LegalEntitlementsCalculator()
deadlines_tracker   = LegalDeadlinesTracker()
slip_detector       = LegalSlipDetector()
case_analyzer       = CaseStrengthAnalyzer()
info_extractor      = LegalInfoExtractor()
