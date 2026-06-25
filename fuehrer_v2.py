"""
Führer 
سري | سياق طويل | حفظ دائم | أي نموذج ذكاء
"""

import streamlit as st
import re, io, os, json, logging, hashlib, base64
from datetime import datetime, timedelta
from typing import Dict, List
import urllib.request, urllib.error

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("fuehrer")

st.set_page_config(
    page_title="⚖️ Führer",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════
# خلفية الصورة — مُضمَّنة مباشرة
# ══════════════════════════════════════════════
BG_B64 = "iVBORw0KGgoAAAANSUhEUgAAAuMAAAPxCAYAAABQFKOA"  # placeholder

def set_bg(b64: str):
    st.markdown(f"""
<style>
.stApp {{
    background-image: url("data:image/png;base64,{b64}");
    background-size: cover;
    background-position: center;
    background-attachment: fixed;
    background-repeat: no-repeat;
}}
.stApp::before {{
    content: '';
    position: fixed;
    top: 0; left: 0; right: 0; bottom: 0;
    background: rgba(8, 12, 20, 0.82);
    z-index: 0;
    pointer-events: none;
}}
</style>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# CSS الكامل
# ══════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;700;900&display=swap');
*{box-sizing:border-box}
.stApp{color:#e8e0d0;font-family:'Cairo',sans-serif;direction:rtl;background:#080c14}
[data-testid="stSidebar"]{background:rgba(13,19,32,0.97)!important;border-left:1px solid #1e2a40;backdrop-filter:blur(10px)}
[data-testid="stSidebar"] *{color:#c8c0b0!important}
h1,h2,h3{color:#f0c040!important;font-weight:700}
/* Tabs */
.stTabs [data-baseweb="tab-list"]{background:rgba(13,19,32,0.9);border-bottom:2px solid #1e2a40;gap:3px;padding:4px;border-radius:8px 8px 0 0;backdrop-filter:blur(8px)}
.stTabs [data-baseweb="tab"]{background:transparent!important;color:#8090a0!important;border:1px solid transparent!important;border-radius:6px!important;padding:7px 14px!important;font-size:13px;font-family:'Cairo',sans-serif}
.stTabs [data-baseweb="tab"][aria-selected="true"]{background:#1a2235!important;color:#f0c040!important;border-color:#f0c040!important;font-weight:700}
.stTabs [data-baseweb="tab-panel"]{background:rgba(10,15,26,0.92);border:1px solid #1e2a40;border-radius:0 0 8px 8px;padding:18px;backdrop-filter:blur(8px)}
/* Inputs */
.stTextInput>div>div>input,.stTextArea textarea{background:rgba(13,19,32,0.95)!important;color:#e8e0d0!important;border:1px solid #2a3a55!important;border-radius:6px!important;font-family:'Cairo',sans-serif!important}
.stTextInput>div>div>input:focus,.stTextArea textarea:focus{border-color:#f0c040!important;box-shadow:0 0 0 2px rgba(240,192,64,.15)!important}
/* Buttons */
.stButton>button{background:linear-gradient(135deg,#c8a020,#f0c040)!important;color:#0a0f1a!important;border:none!important;border-radius:6px!important;font-weight:700!important;font-family:'Cairo',sans-serif!important;padding:10px 18px!important;transition:all .2s!important}
.stButton>button:hover{transform:translateY(-1px);box-shadow:0 4px 20px rgba(240,192,64,.4)!important}
/* Metrics */
[data-testid="stMetric"]{background:rgba(13,19,32,0.9);border:1px solid #1e2a40;border-radius:8px;padding:12px 16px;backdrop-filter:blur(6px)}
[data-testid="stMetricLabel"]{color:#8090a0!important;font-size:12px}
[data-testid="stMetricValue"]{color:#f0c040!important;font-weight:700;font-size:22px}
/* Select */
.stSelectbox [data-baseweb="select"]>div{background:rgba(13,19,32,0.95)!important;border-color:#2a3a55!important;color:#e8e0d0!important}
/* Chat */
.chat-user{background:rgba(26,34,53,0.95);border:1px solid #2a3a55;border-radius:12px 12px 2px 12px;padding:12px 16px;margin:8px 0;max-width:82%;float:right;clear:both;direction:rtl;backdrop-filter:blur(4px)}
.chat-ai{background:rgba(13,26,42,0.95);border:1px solid #1e3a50;border-radius:12px 12px 12px 2px;padding:12px 16px;margin:8px 0;max-width:88%;float:left;clear:both;direction:rtl;border-left:3px solid #f0c040;backdrop-filter:blur(4px)}
.chat-wrap{overflow:hidden;min-height:60px}
/* Cards */
.mem-card{background:rgba(13,19,32,0.92);border:1px solid #1e2a40;border-radius:8px;padding:12px;margin:5px 0;direction:rtl;backdrop-filter:blur(4px)}
.mem-card:hover{border-color:#f0c040}
.ok-card{background:rgba(40,100,60,.2);border:1px solid rgba(64,192,96,.4);border-radius:6px;padding:9px 14px;margin:3px 0;direction:rtl}
.bad-card{background:rgba(100,30,30,.2);border:1px solid rgba(192,64,64,.4);border-radius:6px;padding:9px 14px;margin:3px 0;direction:rtl}
.rule-card{background:rgba(13,26,42,0.92);border-right:4px solid #f0c040;border-radius:0 6px 6px 0;padding:9px 14px;margin:3px 0;direction:rtl;font-size:14px}
.tl-item{border-right:2px solid #2a3a55;padding:8px 16px 8px 0;margin:7px 0;position:relative;direction:rtl}
.tl-item::before{content:'';width:10px;height:10px;background:#f0c040;border-radius:50%;position:absolute;right:-6px;top:12px}
.badge{display:inline-block;background:rgba(26,34,53,0.95);border:1px solid #f0c040;color:#f0c040;border-radius:4px;padding:2px 8px;font-size:11px;font-weight:600;margin:2px}
.hdr{background:linear-gradient(135deg,rgba(13,19,32,0.95),rgba(26,34,53,0.95));border:1px solid #1e2a40;border-bottom:2px solid #f0c040;border-radius:8px;padding:18px 24px;margin-bottom:16px;direction:rtl;backdrop-filter:blur(10px)}
.law-card{background:rgba(13,19,32,0.92);border:1px solid #1e2a40;border-radius:8px;padding:12px;margin:5px 0;direction:rtl}
.law-card:hover{border-color:#4080c0}
hr{border-color:#1e2a40!important}
::-webkit-scrollbar{width:5px}
::-webkit-scrollbar-track{background:rgba(8,12,20,0.5)}
::-webkit-scrollbar-thumb{background:#2a3a55;border-radius:3px}
::-webkit-scrollbar-thumb:hover{background:#f0c040}
/* File uploader */
[data-testid="stFileUploader"]{background:rgba(13,19,32,0.9)!important;border:2px dashed #2a3a55!important;border-radius:8px!important}
/* Checkbox */
.stCheckbox label{color:#c8c0b0!important}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# مجلدات التخزين
# ══════════════════════════════════════════════
DATA_DIR     = "fuehrer_data"
SESSIONS_DIR = os.path.join(DATA_DIR, "sessions")
MEMORY_FILE  = os.path.join(DATA_DIR, "memory.json")
SETTINGS_FILE= os.path.join(DATA_DIR, "settings.json")
LAW_FILE     = os.path.join(DATA_DIR, "law_db.json")
BG_FILE      = os.path.join(DATA_DIR, "bg.b64")

for d in [DATA_DIR, SESSIONS_DIR]:
    os.makedirs(d, exist_ok=True)

# ══════════════════════════════════════════════
# دوال التخزين
# ══════════════════════════════════════════════
def load_json(path, default):
    try:
        if os.path.exists(path):
            with open(path,"r",encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return default

def save_json(path, data):
    try:
        with open(path,"w",encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logger.error("save %s: %s", path, e)
        return False

def list_sessions():
    sessions = []
    try:
        for fname in sorted(os.listdir(SESSIONS_DIR), reverse=True):
            if fname.endswith(".json"):
                data = load_json(os.path.join(SESSIONS_DIR, fname), {})
                sessions.append({
                    "id":      fname.replace(".json",""),
                    "name":    data.get("name","جلسة"),
                    "count":   len(data.get("messages",[])),
                    "updated": data.get("updated",""),
                })
    except Exception:
        pass
    return sessions

def load_session(sid):
    return load_json(os.path.join(SESSIONS_DIR,f"{sid}.json"),
                     {"name":"جلسة جديدة","messages":[],"updated":""})

def save_session(sid, data):
    data["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    return save_json(os.path.join(SESSIONS_DIR,f"{sid}.json"), data)

def delete_session(sid):
    path = os.path.join(SESSIONS_DIR,f"{sid}.json")
    try:
        if os.path.exists(path): os.remove(path)
    except Exception: pass

def new_sid():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

# ══════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════
_saved = load_json(SETTINGS_FILE, {})

def _init():
    defs = {
        "memory":       load_json(MEMORY_FILE, []),
        "law_db":       load_json(LAW_FILE, []),
        "docs":         [],
        "pending_q":    "",
        "current_sid":  None,
        "current_msgs": [],
        # إعدادات النموذج
        "ai_provider":  _saved.get("ai_provider","Gemini"),
        "ai_endpoint":  _saved.get("ai_endpoint","https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"),
        "ai_key":       _saved.get("ai_key",""),
        "ai_model":     _saved.get("ai_model","gemini-2.0-flash"),
        "ai_format":    _saved.get("ai_format","gemini"),
        "case_type":    "قضية عمالية",
        "bg_b64":       "",
    }
    for k,v in defs.items():
        if k not in st.session_state:
            st.session_state[k] = v
    # تحميل الخلفية المحفوظة
    if not st.session_state.bg_b64 and os.path.exists(BG_FILE):
        with open(BG_FILE,"r") as f:
            st.session_state.bg_b64 = f.read().strip()

_init()

def save_settings():
    save_json(SETTINGS_FILE, {
        "ai_provider": st.session_state.ai_provider,
        "ai_endpoint": st.session_state.ai_endpoint,
        "ai_key":      st.session_state.ai_key,
        "ai_model":    st.session_state.ai_model,
        "ai_format":   st.session_state.ai_format,
    })

def save_memory():
    save_json(MEMORY_FILE, st.session_state.memory)

def save_law():
    save_json(LAW_FILE, st.session_state.law_db)

# تطبيق الخلفية
if st.session_state.bg_b64:
    set_bg(st.session_state.bg_b64)

# ══════════════════════════════════════════════
# استخراج القوانين من الملفات النصية
# ══════════════════════════════════════════════
def extract_laws_from_text(text: str, source_name: str = "") -> List[Dict]:
    records = []
    parts = re.split(
        r'(?=المادة\s+(?:الأولى|الثانية|الثالثة|الرابعة|الخامسة|السادسة|السابعة|الثامنة|التاسعة|العاشرة|[\u0600-\u06ff]{3,20})\s*[:\n])',
        text)
    current_law = "الأنظمة السعودية"
    for part in parts:
        part = part.strip()
        if len(part) < 40: continue
        law_m = re.search(r'نظام\s+[\u0600-\u06ff\s]{4,35}(?=\n|،|\.)', part)
        if law_m: current_law = law_m.group(0).strip()
        art_m = re.match(r'(المادة\s+[\u0600-\u06ff\s\d]{2,40}?)(?::|[\n])', part)
        article = art_m.group(1).strip() if art_m else ""
        clean = re.sub(r'https?://\S+', '', part)
        clean = re.sub(r'\s+', ' ', clean).strip()
        arabic_n = sum(1 for c in clean if '\u0600'<=c<='\u06ff')
        if arabic_n < 20 or len(clean) < 40: continue
        records.append({
            "text":     clean[:800],
            "article":  article,
            "law_name": current_law,
            "source":   source_name,
            "ts":       datetime.now().strftime("%Y-%m-%d"),
        })
    return records

def extract_laws_from_pdf(raw: bytes, source_name: str = "") -> List[Dict]:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip(): text += t + "\n"
    except Exception:
        try:
            import PyPDF2
            for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                t = pg.extract_text() or ""
                if t.strip(): text += t + "\n"
        except Exception: pass
    return extract_laws_from_text(text, source_name) if text else []

def extract_laws_from_docx(raw: bytes, source_name: str = "") -> List[Dict]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return extract_laws_from_text(text, source_name)
    except Exception:
        return []

# ══════════════════════════════════════════════
# DOCUMENT INTELLIGENCE — كل أنواع الملفات
# ══════════════════════════════════════════════
def _bytes(f):
    if hasattr(f,"getvalue"): return f.getvalue()
    try:
        p=f.tell(); d=f.read(); f.seek(p); return d
    except Exception: return f.read()

def _norm(t):
    return re.sub(r"\s+"," ", t or "").strip()

class DocIntel:
    SUPPORTED = ["pdf","docx","txt","md","json","csv","rtf","xml","html","htm"]

    def extract(self, f) -> str:
        ext = (getattr(f,"name","") or "").rsplit(".",1)[-1].lower()
        raw = _bytes(f)
        try:
            if ext == "pdf":              return self._pdf(raw)
            if ext == "docx":             return self._docx(raw)
            if ext in ("txt","md","rtf"): return _norm(raw.decode("utf-8",errors="ignore"))
            if ext in ("html","htm"):
                text = re.sub(r'<[^>]+>','',raw.decode("utf-8",errors="ignore"))
                return _norm(text)
            if ext == "xml":
                text = re.sub(r'<[^>]+>','',raw.decode("utf-8",errors="ignore"))
                return _norm(text)
            if ext == "json":
                return _norm(json.dumps(
                    json.loads(raw.decode("utf-8",errors="ignore")),
                    ensure_ascii=False))
            if ext == "csv":
                import csv
                rows = list(csv.reader(io.StringIO(raw.decode("utf-8",errors="ignore"))))
                return _norm("\n".join(" | ".join(r) for r in rows))
            # أي ملف آخر — محاولة قراءة كنص
            return _norm(raw.decode("utf-8",errors="ignore"))
        except Exception:
            return ""

    def _pdf(self, raw):
        parts = []
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for pg in pdf.pages:
                    t = pg.extract_text() or ""
                    if t.strip(): parts.append(t)
        except Exception:
            try:
                import PyPDF2
                for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                    t = pg.extract_text() or ""
                    if t.strip(): parts.append(t)
            except Exception: pass
        return _norm("\n".join(parts))

    def _docx(self, raw):
        try:
            from docx import Document
            return _norm("\n".join(
                p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text))
        except Exception: return ""

    def entities(self, t):
        return {
            "parties":  list(set(re.findall(
                r"(?:المدعي|المدعى عليه|الشركة|المؤسسة|الموظف|الهيئة)", t or ""))),
            "amounts":  re.findall(r"[\d,]+\s*(?:ريال|درهم|دولار)", t or ""),
            "articles": re.findall(r"المادة\s*[\u0600-\u06FF\d]+", t or ""),
            "dates":    re.findall(r"\d{1,2}/\d{1,2}/\d{2,4}", t or ""),
        }

# ══════════════════════════════════════════════
# RULE ENGINE — 45 قاعدة
# ══════════════════════════════════════════════
RULES = [
    {"c":"days_abandoned>30","o":"⚠️ انقطاع >30 يوم (ترك العمل)","cat":"عمل"},
    {"c":"days_abandoned>15 and days_abandoned<=30","o":"⚠️ انقطاع 15-30 يوم (إنذار)","cat":"عمل"},
    {"c":"days_since_firing>365","o":"⛔ مضى >سنة على الفصل (سقط حق التقاضي)","cat":"تقادم"},
    {"c":"days_since_firing>180 and days_since_firing<=365","o":"⏳ مضى >6 أشهر (تقادم جزئي)","cat":"تقادم"},
    {"c":"no_investigation","o":"⚖️ فصل بلا تحقيق (بطلان)","cat":"إجراءات"},
    {"c":"arbitrary_dismissal","o":"⚖️ فصل تعسفي (تعويض واجب)","cat":"عمل"},
    {"c":"salary_delay","o":"⚖️ تأخير الراتب (تعويض)","cat":"عمل"},
    {"c":"eosb_not_paid","o":"⚖️ مكافأة نهاية الخدمة لم تُصرف","cat":"عمل"},
    {"c":"unlawful_deduction","o":"⚖️ خصم بغير حق (يُرد)","cat":"عمل"},
    {"c":"absence_days>30","o":"⚠️ غياب >30 يوم (فصل)","cat":"غياب"},
    {"c":"absence_days>20 and absence_days<=30","o":"⚠️ غياب 20-30 يوم (إنذار ثانٍ)","cat":"غياب"},
    {"c":"absence_days>15 and absence_days<=20","o":"⚠️ غياب 15-20 يوم (إنذار أول)","cat":"غياب"},
    {"c":"service_length<2","o":"📌 خدمة <2 سنة (نصف شهر/سنة)","cat":"مكافأة"},
    {"c":"service_length>=2 and service_length<5","o":"📌 خدمة 2-5 سنوات (شهر/سنة)","cat":"مكافأة"},
    {"c":"service_length>=5","o":"📌 خدمة ≥5 سنوات (شهر ونصف/سنة)","cat":"مكافأة"},
    {"c":"notification_late","o":"⚖️ تبليغ بعد 7 أيام (إخلال إجرائي)","cat":"إجراءات"},
    {"c":"violation_date_missing","o":"⚖️ تاريخ المخالفة مجهول (لصالحك)","cat":"إجراءات"},
    {"c":"penalty_after_1_year","o":"⛔ سنة على المخالفة بلا عقوبة (سقط)","cat":"تقادم"},
    {"c":"judgment_without_hearing","o":"⚖️ حكم دون سماعك (بطلان)","cat":"إجراءات"},
    {"c":"no_response_90_days","o":"⚖️ 90 يوم بلا رد (موافقة ضمنية)","cat":"إجراءات"},
    {"c":"doc_unsigned","o":"⚖️ مستند غير موقع (لا حجية)","cat":"مستندات"},
    {"c":"forgery_proven","o":"🚨 تزوير مثبت (جريمة جنائية)","cat":"مستندات"},
    {"c":"opponent_hides_doc","o":"⚖️ الخصم يخفي مستنداً (يُحكم ضده)","cat":"مستندات"},
    {"c":"new_evidence_late","o":"📌 أدلة جديدة بعد الميعاد (مقبولة)","cat":"مستندات"},
    {"c":"settlement_offer is True and risk_score>60","o":"🤝 الصلح أفضل من التقاضي","cat":"صلح"},
    {"c":"settlement_offer is True and risk_score<=40","o":"⚖️ الصلح ممكن والقضية قوية","cat":"صلح"},
    {"c":"settlement_broken","o":"⚖️ نقض الصلح (تعويض واجب)","cat":"صلح"},
    {"c":"offer_rejected_by_opponent","o":"📌 رفض الخصم الصلح (تعويض لك)","cat":"صلح"},
    {"c":"reply_delay>30","o":"⏳ تأخير إداري >30 يوم (تعنت)","cat":"تأخير"},
    {"c":"ambiguous_count>3","o":"🔍 عبارات غامضة (طعن محتمل)","cat":"لغوي"},
    {"c":"contradictions>1","o":"⚡ تناقض في مراسلات الخصم","cat":"تناقضات"},
    {"c":"force_majeure is True and days_abandoned>60","o":"📌 عذر قاهر يبرر الانقطاع","cat":"أعذار"},
    {"c":"proven_illness","o":"📌 مرض مثبت (عذر مقبول)","cat":"أعذار"},
    {"c":"natural_disaster","o":"📌 كارثة طبيعية (قوة قاهرة)","cat":"أعذار"},
    {"c":"travel_ban","o":"📌 منع السفر (قوة قاهرة)","cat":"أعذار"},
    {"c":"death_of_relative","o":"📌 وفاة قريب (إجازة رسمية)","cat":"أعذار"},
    {"c":"disproportionate_fine","o":"⚖️ غرامة غير متناسبة (تُخفَّض)","cat":"غرامات"},
    {"c":"fine_illegal","o":"⚖️ غرامة مخالفة للنظام (تُلغى)","cat":"غرامات"},
    {"c":"supreme_court_ruling","o":"⭐ سابقة من المحكمة العليا","cat":"سوابق"},
    {"c":"expert_request_denied","o":"⚖️ رفض الخبرة (إخلال بحق الدفاع)","cat":"إجراءات"},
    {"c":"non_judicial_acknowledgment","o":"📌 إقرار غير قضائي (حجة على المُقِر)","cat":"مستندات"},
    {"c":"opponent_threatens","o":"⚖️ تهديد متكرر (تعسف)","cat":"سلوك"},
    {"c":"witnesses_conflict","o":"⚖️ تناقض الشهود (تُرجح الأعدل)","cat":"شهادات"},
    {"c":"two_vs_one_witness","o":"📌 شاهدان ضد واحد (مقبول)","cat":"شهادات"},
    {"c":"no_registered_letter","o":"⚖️ تبليغ بغير بريد مسجل","cat":"إجراءات"},
]

def eval_rule(cond, ctx):
    try:
        for part in [p.strip() for p in cond.split(" and ")]:
            if not part: continue
            m = re.match(r"^(\w+)\s+is\s+(True|False)$", part)
            if m:
                if bool(ctx.get(m[1],False))!=(m[2]=="True"): return False
                continue
            m = re.match(r"^(\w+)$", part)
            if m:
                if not bool(ctx.get(m[1],False)): return False
                continue
            m = re.match(r"^(\w+)\s*(>=|<=|>|<)\s*([0-9.]+)$", part)
            if m:
                lhs=float(ctx.get(m[1],0)); rhs=float(m[3])
                if not {">":lhs>rhs,">=":lhs>=rhs,"<":lhs<rhs,"<=":lhs<=rhs}[m[2]]: return False
                continue
            m = re.match(r"^(\w+)=='([^']*)'$", part.replace(" ",""))
            if m:
                if str(ctx.get(m[1],""))!=m[2]: return False
                continue
            return False
        return True
    except Exception: return False

def apply_rules(ctx):
    return [{"text":r["o"],"cat":r["cat"]} for r in RULES if eval_rule(r["c"],ctx)]

# ══════════════════════════════════════════════
# TIMELINE
# ══════════════════════════════════════════════
def build_timeline(texts):
    evs = []
    for txt in texts:
        for d in re.findall(r"\d{1,2}/\d{1,2}/\d{2,4}", txt or ""):
            for fmt in ["%d/%m/%Y","%d/%m/%y"]:
                try:
                    evs.append({"date":datetime.strptime(d,fmt),"text":(txt or "")[:200]})
                    break
                except ValueError: pass
    return sorted(evs, key=lambda x: x["date"])

def calc_gaps(evs):
    out=[]
    for i in range(len(evs)-1):
        diff=(evs[i+1]["date"]-evs[i]["date"]).days
        if diff>30:
            out.append({"from":evs[i]["date"].strftime("%d/%m/%Y"),
                        "to":evs[i+1]["date"].strftime("%d/%m/%Y"),"days":diff})
    return out

# ══════════════════════════════════════════════
# MEMORY
# ══════════════════════════════════════════════
def mem_add(text, tags=None, cat="عام"):
    m={"id":hashlib.md5(f"{text}{datetime.now().isoformat()}".encode()).hexdigest()[:8],
       "text":text,"tags":tags or [],"category":cat,
       "ts":datetime.now().strftime("%Y-%m-%d %H:%M")}
    st.session_state.memory.append(m)
    save_memory(); return m["id"]

def mem_del(mid):
    st.session_state.memory=[m for m in st.session_state.memory if m["id"]!=mid]
    save_memory()

def mem_edit(mid, new_text):
    for m in st.session_state.memory:
        if m["id"]==mid:
            m["text"]=new_text
            m["ts"]=datetime.now().strftime("%Y-%m-%d %H:%M")
            break
    save_memory()

# ══════════════════════════════════════════════
# AI — أي نموذج / أي endpoint
# ══════════════════════════════════════════════
def build_system():
    mem_ctx=""
    if st.session_state.memory:
        mem_ctx="\n\nالذاكرة:\n"+"\n".join(
            f"- {m['text'][:150]}" for m in st.session_state.memory[-20:])
    law_ctx=""
    if st.session_state.law_db and st.session_state.current_msgs:
        last_q = next((m["content"] for m in reversed(st.session_state.current_msgs)
                      if m["role"]=="user"), "")
        if last_q:
            q_words = set(re.findall(r"[\u0600-\u06ff]{3,}", last_q))
            scored = sorted(
                [(sum(1 for w in q_words if w in r.get("text","")),r)
                 for r in st.session_state.law_db], reverse=True)
            relevant = [(sc,r) for sc,r in scored if sc>0][:5]
            if relevant:
                law_ctx="\n\nمواد قانونية ذات صلة:\n"
                for sc,r in relevant:
                    law_ctx+=f"• [{r['law_name']}] {r['article']}: {r['text'][:300]}\n"
    doc_ctx=""
    if st.session_state.docs:
        doc_ctx=f"\n\nالمستندات:\n{chr(10).join(st.session_state.docs[:3])[:4000]}"
    return f"""أنت محامٍ ومستشار قانوني سعودي خبير.
تخصصك: نظام العمل، المرافعات الشرعية، الأنظمة السعودية.
- استند للأنظمة السعودية واذكر المواد بالاسم
- كن محدداً وعملياً في إجاباتك
- أجب بالعربية الفصحى الواضحة{mem_ctx}{law_ctx}{doc_ctx}"""

def call_ai(prompt: str) -> str:
    key      = st.session_state.ai_key
    endpoint = st.session_state.ai_endpoint
    model    = st.session_state.ai_model
    fmt      = st.session_state.ai_format
    msgs     = st.session_state.current_msgs

    if not key:    return "❌ أدخل API Key في الإعدادات"
    if not endpoint: return "❌ أدخل رابط الـ API"

    system = build_system()

    # ── Gemini format ──────────────────────
    if fmt == "gemini":
        contents = [{"role":"user","parts":[{"text":system+"\n\nالسؤال: "+prompt}]}]
        for m in msgs[-30:]:
            role = "user" if m["role"]=="user" else "model"
            contents.append({"role":role,"parts":[{"text":m["content"]}]})
        payload = json.dumps({"contents":contents}).encode()
        url = f"{endpoint}?key={key}"
        req = urllib.request.Request(url, data=payload,
              headers={"Content-Type":"application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=90) as r:
                d = json.loads(r.read().decode())
                return d["candidates"][0]["content"]["parts"][0]["text"]
        except urllib.error.HTTPError as e:
            err = e.read().decode()[:300]
            return f"❌ Gemini {e.code}: {err}"
        except Exception as e:
            return f"❌ {e}"

    # ── Anthropic format ───────────────────
    elif fmt == "anthropic":
        messages = []
        for m in msgs[-30:]:
            messages.append({"role":m["role"],"content":m["content"]})
        messages.append({"role":"user","content":prompt})
        payload = json.dumps({
            "model": model,
            "max_tokens": 2048,
            "system": system,
            "messages": messages,
        }).encode()
        req = urllib.request.Request(endpoint, data=payload,
              headers={"Content-Type":"application/json",
                       "x-api-key":key,
                       "anthropic-version":"2023-06-01"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=90) as r:
                d = json.loads(r.read().decode())
                return d["content"][0]["text"]
        except urllib.error.HTTPError as e:
            return f"❌ Claude {e.code}: {e.read().decode()[:200]}"
        except Exception as e:
            return f"❌ {e}"

    # ── OpenAI-compatible format (Groq, Together, Ollama, etc.) ──
    else:
        messages = [{"role":"system","content":system}]
        for m in msgs[-30:]:
            messages.append({"role":m["role"],"content":m["content"]})
        messages.append({"role":"user","content":prompt})
        payload = json.dumps({
            "model":model,
            "messages":messages,
            "max_tokens":2048,
        }).encode()
        req = urllib.request.Request(endpoint, data=payload,
              headers={"Content-Type":"application/json",
                       "Authorization":f"Bearer {key}"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=90) as r:
                d = json.loads(r.read().decode())
                return d["choices"][0]["message"]["content"]
        except urllib.error.HTTPError as e:
            return f"❌ {e.code}: {e.read().decode()[:200]}"
        except Exception as e:
            return f"❌ {e}"

# ══════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════
with st.sidebar:
    st.markdown("### ⚖️ Führer v2.0")
    st.markdown("---")

    # ── إعدادات النموذج ──────────────────────
    with st.expander("🤖 إعدادات النموذج", expanded=True):
        preset = st.selectbox("اختر سريع أو خصّص", [
            "Gemini 2.0 Flash — مجاني",
            "Groq LLaMA 3.3 — مجاني",
            "Claude Sonnet",
            "OpenAI GPT-4o",
            "Ollama محلي",
            "⚙️ مخصص — أدخل يدوياً",
        ], label_visibility="collapsed")

        PRESETS = {
            "Gemini 2.0 Flash — مجاني": {
                "endpoint": "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent",
                "model": "gemini-2.0-flash",
                "format": "gemini",
                "placeholder": "AIza...",
                "link": "https://aistudio.google.com/apikey",
            },
            "Groq LLaMA 3.3 — مجاني": {
                "endpoint": "https://api.groq.com/openai/v1/chat/completions",
                "model": "llama-3.3-70b-versatile",
                "format": "openai",
                "placeholder": "gsk_...",
                "link": "https://console.groq.com",
            },
            "Claude Sonnet": {
                "endpoint": "https://api.anthropic.com/v1/messages",
                "model": "claude-sonnet-4-6",
                "format": "anthropic",
                "placeholder": "sk-ant-...",
                "link": "https://console.anthropic.com",
            },
            "OpenAI GPT-4o": {
                "endpoint": "https://api.openai.com/v1/chat/completions",
                "model": "gpt-4o",
                "format": "openai",
                "placeholder": "sk-...",
                "link": "https://platform.openai.com/api-keys",
            },
            "Ollama محلي": {
                "endpoint": "http://localhost:11434/v1/chat/completions",
                "model": "llama3",
                "format": "openai",
                "placeholder": "ollama",
                "link": "https://ollama.com",
            },
        }

        if preset != "⚙️ مخصص — أدخل يدوياً" and preset in PRESETS:
            p = PRESETS[preset]
            st.session_state.ai_endpoint = p["endpoint"]
            st.session_state.ai_model    = p["model"]
            st.session_state.ai_format   = p["format"]
            st.session_state.ai_provider = preset
            ph = p["placeholder"]
            link = p["link"]
        else:
            ph = "API Key..."
            link = ""
            st.session_state.ai_endpoint = st.text_input(
                "رابط API", value=st.session_state.ai_endpoint,
                placeholder="https://api.example.com/v1/chat/completions")
            st.session_state.ai_model = st.text_input(
                "اسم النموذج", value=st.session_state.ai_model,
                placeholder="gpt-4o / llama3 / ...")
            fmt_choice = st.selectbox("صيغة API",
                ["openai","gemini","anthropic"])
            st.session_state.ai_format = fmt_choice

        k = st.text_input("🔑 API Key", value=st.session_state.ai_key,
                          type="password", placeholder=ph,
                          label_visibility="collapsed")
        if k != st.session_state.ai_key:
            st.session_state.ai_key = k
            save_settings()
        if link:
            st.markdown(f"[احصل على Key مجاني ←]({link})")
        if st.session_state.ai_key:
            st.success("✅ محفوظ")
        if st.button("💾 حفظ الإعدادات", use_container_width=True):
            save_settings()
            st.success("✅ تم الحفظ")

    st.markdown("---")

    # ── الجلسات ──────────────────────────────
    st.markdown("**💬 الجلسات**")
    if st.button("➕ جلسة جديدة", use_container_width=True):
        sid = new_sid()
        st.session_state.current_sid  = sid
        st.session_state.current_msgs = []
        save_session(sid, {"name":"جلسة جديدة","messages":[]})
        st.rerun()

    for s in list_sessions()[:12]:
        c1,c2 = st.columns([5,1])
        with c1:
            active = "🟢 " if s["id"]==st.session_state.current_sid else ""
            lbl = f"{active}{s['name'][:18]} ({s['count']})"
            if st.button(lbl, key=f"s_{s['id']}", use_container_width=True):
                data = load_session(s["id"])
                st.session_state.current_sid  = s["id"]
                st.session_state.current_msgs = data.get("messages",[])
                st.rerun()
        with c2:
            if st.button("🗑", key=f"ds_{s['id']}"):
                delete_session(s["id"])
                if st.session_state.current_sid==s["id"]:
                    st.session_state.current_sid=None
                    st.session_state.current_msgs=[]
                st.rerun()

    st.markdown("---")
    st.markdown("**📋 نوع القضية**")
    st.session_state.case_type = st.selectbox("النوع",[
        "قضية عمالية","نزاع تجاري","قضية عقارية",
        "نزاع إداري","قضية جنائية","إفلاس وتصفية",
    ], label_visibility="collapsed")

    st.markdown("---")
    c1,c2 = st.columns(2)
    with c1: st.metric("الذاكرة", len(st.session_state.memory))
    with c2: st.metric("الجلسات", len(list_sessions()))
    st.metric("مواد القانون", len(st.session_state.law_db))

# ══════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════
st.markdown(f"""
<div class="hdr">
<h1 style="margin:0;font-size:24px">⚖️ Führer  |</h1>
<p style="color:#8090a0;margin:4px 0 0;font-size:12px">
سري تماماً • سياق طويل • حفظ دائم • {st.session_state.ai_provider}
• {len(st.session_state.law_db):,} مادة قانونية
</p>
</div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════
tabs = st.tabs(["🤖 المستشار","📂 الملفات","📅 الجدول","⚖️ التحليل",
                "📜 القواعد","📄 التقارير","🧠 الذاكرة","📚 القانون","⚙️ الإعدادات"])
t_ai,t_files,t_tl,t_analysis,t_rules,t_reports,t_mem,t_law,t_settings = tabs

# ── TAB 1: المستشار ─────────────────────────
with t_ai:
    st.subheader(f"🤖 {st.session_state.ai_provider}")

    if not st.session_state.current_sid:
        st.markdown("""
<div style="text-align:center;padding:40px;color:#8090a0">
<h2>👈 ابدأ بجلسة جديدة</h2>
<p>اضغط "جلسة جديدة" من الشريط الجانبي</p>
</div>""", unsafe_allow_html=True)
    else:
        sess = load_session(st.session_state.current_sid)
        new_name = st.text_input("📝 اسم الجلسة",
                                  value=sess.get("name","جلسة"),
                                  key="sess_name_inp")
        if new_name != sess.get("name",""):
            sess["name"] = new_name
            sess["messages"] = st.session_state.current_msgs
            save_session(st.session_state.current_sid, sess)

        # Quick prompts
        cols = st.columns(4)
        for i,(col,q) in enumerate(zip(cols,[
            "حلل وضعي القانوني",
            "ما نقاط قوتي؟",
            "ما المواعيد النظامية؟",
            "اقترح استراتيجية",
        ])):
            with col:
                if st.button(q, key=f"qp{i}", use_container_width=True):
                    st.session_state.pending_q = q

        st.markdown("---")

        # عرض المحادثة
        st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
        for msg in st.session_state.current_msgs:
            cls = "chat-user" if msg["role"]=="user" else "chat-ai"
            ico = "👤" if msg["role"]=="user" else "⚖️"
            content = msg["content"].replace("\n","<br>")
            ts = msg.get("ts","")
            st.markdown(
                f'<div class="{cls}">{ico} {content}'
                f'<br><small style="color:#556;font-size:10px">⏱ {ts}</small></div>',
                unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("---")

        user_inp = st.text_area(
            "سؤالك",
            value=st.session_state.pending_q,
            height=100,
            placeholder="مثال: تأخر راتبي 3 أشهر وأُشعرت بالفصل — ما حقوقي القانونية؟",
            key="chat_inp",
        )

        bc1,bc2,bc3 = st.columns([3,1,1])
        with bc1: send_btn = st.button("📤 إرسال", use_container_width=True)
        with bc2:
            if st.button("🗑️ مسح", use_container_width=True):
                st.session_state.current_msgs=[]
                sess["messages"]=[]
                save_session(st.session_state.current_sid, sess)
                st.rerun()
        with bc3:
            if st.button("💾 حفظ", use_container_width=True):
                sess["messages"]=st.session_state.current_msgs
                save_session(st.session_state.current_sid, sess)
                st.success("✅")

        if send_btn and user_inp.strip():
            st.session_state.pending_q=""
            ts=datetime.now().strftime("%H:%M")
            st.session_state.current_msgs.append(
                {"role":"user","content":user_inp,"ts":ts})
            with st.spinner("⚖️ يحلل..."):
                resp=call_ai(user_inp)
            st.session_state.current_msgs.append(
                {"role":"assistant","content":resp,"ts":ts})
            sess["messages"]=st.session_state.current_msgs
            save_session(st.session_state.current_sid, sess)
            if len(resp)>80 and "❌" not in resp:
                mem_add(f"س: {user_inp[:80]} | ج: {resp[:150]}...",
                       tags=["محادثة",st.session_state.case_type],cat="محادثة")
            st.rerun()

        if st.session_state.current_msgs:
            chat_txt="\n\n".join(
                f"{'أنت' if m['role']=='user' else 'المستشار'} [{m.get('ts','')}]:\n{m['content']}"
                for m in st.session_state.current_msgs)
            st.download_button("⬇️ تحميل المحادثة",
                data=chat_txt.encode("utf-8"),
                file_name=f"محادثة_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain")

# ── TAB 2: الملفات ──────────────────────────
with t_files:
    st.subheader("📂 رفع وتحليل المستندات")
    st.markdown('<small style="color:#8090a0">يدعم: PDF · DOCX · TXT · MD · JSON · CSV · HTML · XML · وأي ملف نصي</small>', unsafe_allow_html=True)

    uploaded = st.file_uploader(
        "اختر الملفات",
        type=None,  # كل الأنواع
        accept_multiple_files=True,
        label_visibility="collapsed",
    )
    if uploaded:
        st.info(f"✅ {len(uploaded)} ملف")
        c1,c2 = st.columns(2)
        with c1:
            if st.button("🔍 تحليل وعرض", use_container_width=True):
                di=DocIntel()
                texts=[]
                for f in uploaded:
                    with st.expander(f"📄 {f.name}"):
                        txt=di.extract(f)
                        if txt:
                            texts.append(txt)
                            ents=di.entities(txt)
                            st.text(txt[:500]+("..." if len(txt)>500 else ""))
                            if ents["articles"]:
                                st.markdown("**المواد:** "+" ".join(
                                    f'<span class="badge">{a}</span>'
                                    for a in ents["articles"][:6]),
                                    unsafe_allow_html=True)
                            if ents["dates"]:
                                st.markdown(f"**تواريخ:** {', '.join(ents['dates'][:5])}")
                            if ents["amounts"]:
                                st.markdown(f"**مبالغ:** {', '.join(ents['amounts'][:5])}")
                        else:
                            st.warning("⚠️ لم يُستخرج نص")
                st.session_state.docs=texts
                st.success(f"✅ {len(texts)} ملف | {sum(len(t) for t in texts):,} حرف")

        with c2:
            if st.button("📚 إضافة للقاعدة القانونية", use_container_width=True):
                total=0
                for f in uploaded:
                    raw=_bytes(f)
                    ext=(f.name or "").rsplit(".",1)[-1].lower()
                    if ext=="pdf":
                        records=extract_laws_from_pdf(raw, f.name)
                    elif ext=="docx":
                        records=extract_laws_from_docx(raw, f.name)
                    else:
                        text=raw.decode("utf-8",errors="ignore")
                        records=extract_laws_from_text(text, f.name)
                    st.session_state.law_db.extend(records)
                    total+=len(records)
                save_law()
                st.success(f"✅ {total} مادة أُضيفت للقاعدة")

# ── TAB 3: الجدول الزمني ────────────────────
with t_tl:
    st.subheader("📅 الجدول الزمني")
    if not st.session_state.docs:
        st.info("⚠️ ارفع الملفات أولاً")
    else:
        tl=build_timeline(st.session_state.docs)
        gaps=calc_gaps(tl)
        m1,m2=st.columns(2)
        with m1: st.metric("الأحداث",len(tl))
        with m2: st.metric("الفجوات",len(gaps))
        for ev in tl:
            st.markdown(
                f'<div class="tl-item"><strong>{ev["date"].strftime("%d/%m/%Y")}</strong>'
                f'<br><span style="color:#a0b0c0;font-size:13px">{ev["text"][:120]}...</span></div>',
                unsafe_allow_html=True)
        if gaps:
            st.markdown("### ⚠️ الفجوات")
            for g in gaps:
                st.error(f"⏰ {g['days']} يوم — من {g['from']} إلى {g['to']}")

# ── TAB 4: التحليل ──────────────────────────
with t_analysis:
    st.subheader("⚖️ التحليل")
    if not st.session_state.docs:
        st.info("⚠️ ارفع الملفات أولاً")
    else:
        texts=st.session_state.docs
        tl=build_timeline(texts)
        gaps=calc_gaps(tl)
        contras=[]
        for i,t in enumerate(texts):
            dates=re.findall(r"\d{1,2}/\d{1,2}/\d{2,4}",t or "")
            if len(dates)>=2 and dates[0]==dates[1]:
                contras.append(f"تناقض في التواريخ بالملف {i+1}")
        ss=sum(1 for t in texts for k in ["تهديد","فوراً","عاجل"] if k in t)
        risk=min(max(len(gaps)*2+len(contras)*5+ss+(10 if len(tl)<2 else 0),0),100)
        cred=max(100-sum(5 for t in texts if "نحن نؤكد" in t)
                    -sum(10 for t in texts if "مادة" in t and "خطأ" in t),0)
        mc=st.columns(4)
        with mc[0]: st.metric("مستوى الخطر",f"{risk}/100")
        with mc[1]: st.metric("مصداقية الخصم",f"{cred}/100")
        with mc[2]: st.metric("التناقضات",len(contras))
        with mc[3]: st.metric("الفجوات",len(gaps))
        color="#c04040" if risk>70 else "#c08020" if risk>40 else "#40c060"
        st.markdown(
            f'<div style="background:rgba(13,19,32,0.9);border:1px solid #1e2a40;'
            f'border-radius:6px;padding:8px;margin:8px 0">'
            f'<div style="background:{color};width:{risk}%;height:8px;border-radius:4px"></div>'
            f'<small style="color:#8090a0">الخطر: {risk}%</small></div>',
            unsafe_allow_html=True)
        strs,weaks=[],[]
        for ev in tl:
            t=(ev.get("text") or "").lower()
            full=ev.get("text") or ""
            if "أقر" in t or "اعترف" in t: weaks.append("اعتراف ضمني من الخصم")
            if any(k in t for k in ["عذر","مرض","ظروف"]): strs.append("أعذار رسمية موثقة")
            if "المادة" in full: strs.append("استشهاد بمواد نظامية")
            if "تهديد" in t: weaks.append("لغة تهديدية من الخصم")
        strs=list(set(strs)); weaks=list(set(weaks))
        sc1,sc2=st.columns(2)
        with sc1:
            st.markdown("### ✅ نقاط القوة")
            for s in strs: st.markdown(f'<div class="ok-card">✅ {s}</div>',unsafe_allow_html=True)
            if not strs: st.info("لا توجد")
        with sc2:
            st.markdown("### ❌ نقاط الضعف")
            for w in weaks: st.markdown(f'<div class="bad-card">⚠️ {w}</div>',unsafe_allow_html=True)
            if not weaks: st.success("لا توجد")

# ── TAB 5: القواعد ──────────────────────────
with t_rules:
    st.subheader(f"📜 محرك القواعد — {len(RULES)} قاعدة")
    with st.expander("⚙️ بيانات القضية",expanded=True):
        rc1,rc2,rc3=st.columns(3)
        with rc1:
            d_aban =st.number_input("أيام الانقطاع",0,3000,0)
            d_fire =st.number_input("أيام منذ الفصل",0,3000,0)
            d_reply=st.number_input("تأخر رد الخصم",0,365,0)
            d_abs  =st.number_input("أيام الغياب",0,365,0)
        with rc2:
            svc   =st.number_input("سنوات الخدمة",0.0,50.0,0.0,0.5)
            rscore=st.number_input("درجة الخطر",0,100,50)
        with rc3:
            no_inv =st.checkbox("فصل بلا تحقيق")
            arb_dis=st.checkbox("فصل تعسفي")
            fm     =st.checkbox("عذر قاهر")
            settl  =st.checkbox("عرض صلح")
            sal_del=st.checkbox("تأخير الراتب")
            eosb   =st.checkbox("مكافأة لم تُصرف")
            ill    =st.checkbox("مرض مثبت")
            no_resp=st.checkbox("90 يوم بلا رد")
            forgery=st.checkbox("تزوير مثبت")
    if st.button("🔍 تطبيق القواعد",use_container_width=True):
        ctx={
            "days_abandoned":d_aban,"days_since_firing":d_fire,
            "reply_delay":d_reply,"absence_days":d_abs,
            "service_length":svc,"risk_score":rscore,
            "no_investigation":no_inv,"arbitrary_dismissal":arb_dis,
            "force_majeure":fm,"settlement_offer":settl,
            "salary_delay":sal_del,"eosb_not_paid":eosb,
            "proven_illness":ill,"no_response_90_days":no_resp,
            "forgery_proven":forgery,
        }
        alerts=apply_rules(ctx)
        if alerts:
            cats={}
            for a in alerts: cats.setdefault(a["cat"],[]).append(a["text"])
            for cat,items in cats.items():
                st.markdown(f"**{cat}**")
                for item in items:
                    st.markdown(f'<div class="rule-card">{item}</div>',unsafe_allow_html=True)
        else:
            st.success("✅ لا تنبيهات")

# ── TAB 6: التقارير ─────────────────────────
with t_reports:
    st.subheader("📄 التقارير واللوائح")
    rp1,rp2=st.columns(2)
    with rp1:
        st.markdown("### 📊 تقرير شامل")
        if st.button("🖨️ إنشاء",use_container_width=True):
            if not st.session_state.docs:
                st.warning("ارفع الملفات أولاً")
            else:
                texts=st.session_state.docs
                tl=build_timeline(texts)
                gaps=calc_gaps(tl)
                report=f"""تقرير قانوني — {datetime.now().strftime('%d/%m/%Y %H:%M')}
نوع القضية: {st.session_state.case_type}
النموذج المستخدم: {st.session_state.ai_provider}
{'='*40}
الأحداث: {len(tl)} | الفجوات: {len(gaps)}
الذاكرة: {len(st.session_state.memory)} سجل
مواد القانون: {len(st.session_state.law_db)} مادة

الفجوات الزمنية:
"""+"".join(f"• {g['days']} يوم: {g['from']} → {g['to']}\n" for g in gaps)
                st.text_area("التقرير",report,height=300)
                st.download_button("⬇️ تحميل",
                    data=report.encode("utf-8"),
                    file_name=f"تقرير_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain")
    with rp2:
        st.markdown("### ✍️ مسودة اللائحة")
        tmpl=st.selectbox("النوع",["مذكرة دفاع","صحيفة دعوى","عريضة اعتراض","إنذار رسمي"])
        court_n =st.text_input("المحكمة","محكمة العمل")
        case_n  =st.text_input("رقم القضية","___/___/____")
        client_n=st.text_input("الموكل","")
        oppon_n =st.text_input("الخصم","")
        facts_n =st.text_area("الوقائع","",height=100)
        reqs_n  =st.text_area("الطلبات","إلغاء القرار والتعويض",height=100)
        if st.button("✍️ إنشاء",use_container_width=True):
            dn=datetime.now().strftime("%d/%m/%Y")
            drafts={
"مذكرة دفاع":f"""بسم الله الرحمن الرحيم
السيد/ رئيس {court_n} المحترم
الموضوع: مذكرة دفاع — الدعوى ({case_n})
المقدم من: {client_n}  ضد: {oppon_n}

الوقائع:
{facts_n}

الطلبات:
{reqs_n}

{client_n} — {dn}""",
"صحيفة دعوى":f"""بسم الله الرحمن الرحيم
رئيس {court_n} المحترم
المدعي: {client_n} | المدعى عليه: {oppon_n}
الوقائع: {facts_n}
الطلبات: {reqs_n}
{dn}""",
"عريضة اعتراض":f"""بسم الله الرحمن الرحيم
رئيس {court_n} المحترم
اعتراض على القرار ({case_n}) — مقدم من {client_n}
الأسباب: {facts_n}
الطلبات: {reqs_n} — {dn}""",
"إنذار رسمي":f"""بسم الله الرحمن الرحيم
إلى: {oppon_n} | من: {client_n} | {dn}
أُنذركم بشأن: {facts_n}
في حال عدم الاستجابة خلال 15 يوماً ستُتَّخذ الإجراءات القانونية.
{reqs_n}""",
            }
            draft=drafts.get(tmpl,"")
            st.text_area("المسودة",draft,height=400)
            st.download_button("⬇️ تحميل",
                data=draft.encode("utf-8"),
                file_name=f"مسودة_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain")

# ── TAB 7: الذاكرة ──────────────────────────
with t_mem:
    st.subheader("🧠 الذاكرة الدائمة")

    # إضافة من ملف
    with st.expander("📁 إضافة من ملف"):
        mem_file=st.file_uploader("ارفع ملفاً لإضافته للذاكرة",
                                   type=None, key="mem_file_up")
        mem_file_cat=st.selectbox("الفئة",
            ["قضية","موكل","حكم","ملاحظة","قانون","عام"],key="mfc")
        if mem_file and st.button("📥 إضافة الملف للذاكرة"):
            di=DocIntel()
            txt=di.extract(mem_file)
            if txt:
                # قسّم الملف الكبير لأجزاء
                chunks=[txt[i:i+500] for i in range(0,min(len(txt),5000),500)]
                for ch in chunks:
                    mem_add(ch,tags=[mem_file.name],cat=mem_file_cat)
                st.success(f"✅ {len(chunks)} قطعة من {mem_file.name}")
            else:
                st.warning("⚠️ لم يُستخرج نص")

    # إضافة يدوية
    with st.expander("✏️ إضافة يدوية"):
        mt   =st.text_area("النص",height=100,placeholder="مثال: الموكل يعمل منذ 2019")
        mcat =st.selectbox("الفئة",["قضية","موكل","حكم","ملاحظة","استراتيجية","قانون","عام"])
        mtags=st.text_input("وسوم (فاصلة)")
        if st.button("💾 حفظ"):
            if mt.strip():
                tags=[x.strip() for x in mtags.split(",") if x.strip()]
                mid=mem_add(mt,tags,mcat)
                st.success(f"✅ (ID: {mid})")
                st.rerun()

    mq=st.text_input("🔍 بحث في الذاكرة")
    q=mq.lower()
    mems=[m for m in st.session_state.memory
          if not mq or q in m["text"].lower()
          or any(q in t.lower() for t in m.get("tags",[]))]
    st.markdown(f"**{len(mems)} ذاكرة**")
    for m in reversed(mems):
        ec1,ec2,ec3=st.columns([8,1,1])
        with ec1:
            badges="".join(f'<span class="badge">{t}</span>' for t in m.get("tags",[]))
            st.markdown(
                f'<div class="mem-card">'
                f'<small style="color:#8090a0">{m.get("ts","")} · {m.get("category","")}</small>'
                f'<br>{m["text"][:200]}<br>{badges}</div>',
                unsafe_allow_html=True)
        with ec2:
            if st.button("✏️",key=f"e_{m['id']}"):
                st.session_state[f"edit_{m['id']}"]=True
        with ec3:
            if st.button("🗑",key=f"d_{m['id']}"):
                mem_del(m["id"]); st.rerun()
        if st.session_state.get(f"edit_{m['id']}"):
            new_t=st.text_area("تعديل",value=m["text"],key=f"et_{m['id']}",height=100)
            if st.button("✅ حفظ",key=f"sv_{m['id']}"):
                mem_edit(m["id"],new_t)
                del st.session_state[f"edit_{m['id']}"]
                st.rerun()

    st.markdown("---")
    ex1,ex2=st.columns(2)
    with ex1:
        if st.button("📤 تصدير"):
            d=json.dumps(st.session_state.memory,ensure_ascii=False,indent=2)
            st.download_button("⬇️ JSON",d.encode("utf-8"),"memory.json","application/json")
    with ex2:
        mf=st.file_uploader("📥 استيراد JSON",type=["json"],key="mf_up")
        if mf:
            try:
                imported=json.loads(mf.read())
                existing={m["id"] for m in st.session_state.memory}
                new_ones=[m for m in imported if m.get("id") not in existing]
                st.session_state.memory.extend(new_ones)
                save_memory()
                st.success(f"✅ {len(new_ones)} ذاكرة")
            except Exception as e:
                st.error(f"❌ {e}")

# ── TAB 8: القانون ──────────────────────────
with t_law:
    st.subheader("📚 قاعدة الأنظمة السعودية")

    # رفع أي ملف
    st.markdown("### 📁 إضافة قوانين من ملف")
    st.markdown('<small style="color:#8090a0">يدعم أي ملف: TXT · PDF · DOCX · MD · JSON · CSV</small>', unsafe_allow_html=True)
    law_file=st.file_uploader("ارفع ملف القوانين",type=None,key="law_file_up")
    if law_file and st.button("📥 استخراج وإضافة للقاعدة",use_container_width=True):
        raw=_bytes(law_file)
        ext=(law_file.name or "").rsplit(".",1)[-1].lower()
        with st.spinner(f"جاري استخراج {law_file.name}..."):
            if ext=="pdf":
                records=extract_laws_from_pdf(raw,law_file.name)
            elif ext=="docx":
                records=extract_laws_from_docx(raw,law_file.name)
            elif ext=="json":
                try:
                    records=json.loads(raw.decode("utf-8",errors="ignore"))
                    if isinstance(records,list):
                        pass  # استخدمه مباشرة
                    else:
                        records=[]
                except Exception:
                    records=[]
            else:
                text=raw.decode("utf-8",errors="ignore")
                records=extract_laws_from_text(text,law_file.name)

        if records:
            st.session_state.law_db.extend(records)
            save_law()
            st.success(f"✅ {len(records)} مادة من {law_file.name}")
            # إحصائيات
            law_names=list(set(r.get("law_name","") for r in records))
            st.markdown("**الأنظمة المضافة:**")
            for ln in law_names[:8]:
                if ln: st.markdown(f'<span class="badge">{ln}</span>',unsafe_allow_html=True)
        else:
            st.warning("⚠️ لم يُستخرج محتوى — تحقق من الملف")

    st.markdown("---")

    # إضافة يدوية
    with st.expander("✏️ إضافة مادة يدوياً"):
        ma_text=st.text_area("نص المادة",height=100,key="ma_t")
        ma_art =st.text_input("اسم المادة",key="ma_a")
        ma_law =st.text_input("اسم النظام",key="ma_l")
        if st.button("➕ إضافة"):
            if ma_text.strip():
                st.session_state.law_db.append({
                    "text":ma_text,"article":ma_art,
                    "law_name":ma_law or "نظام يدوي",
                    "source":"manual",
                    "ts":datetime.now().strftime("%Y-%m-%d"),
                })
                save_law()
                st.success("✅")

    # إحصائيات وبحث
    if st.session_state.law_db:
        lm1,lm2=st.columns(2)
        with lm1: st.metric("إجمالي المواد",len(st.session_state.law_db))
        with lm2:
            ln_count=len(set(r.get("law_name","") for r in st.session_state.law_db))
            st.metric("عدد الأنظمة",ln_count)

        law_q=st.text_input("🔍 ابحث في الأنظمة")
        if law_q:
            q_words=set(re.findall(r"[\u0600-\u06FF]{3,}",law_q))
            scored=sorted(
                [(sum(1 for w in q_words
                      if w in r.get("text","") or w in r.get("article","")),r)
                 for r in st.session_state.law_db],
                reverse=True)
            results=[(sc,r) for sc,r in scored if sc>0][:10]
            if results:
                st.markdown(f"**{len(results)} نتيجة:**")
                for sc,r in results:
                    with st.expander(f"📜 {r.get('article','')} — {r.get('law_name','')} ({sc})"):
                        st.write(r["text"])
                        bc1,bc2=st.columns(2)
                        with bc1:
                            if st.button("💾 حفظ في الذاكرة",key=f"ls_{hash(r['text'])%99999}"):
                                mem_add(f"[{r.get('law_name','')}] {r.get('article','')}: {r['text'][:200]}",
                                       tags=["قانون",r.get("law_name","")],cat="قانون")
                                st.success("✅")
            else:
                st.info(f"لا نتائج لـ '{law_q}'")

        # تصدير
        col_exp,col_clr=st.columns(2)
        with col_exp:
            if st.button("📤 تصدير القاعدة"):
                d=json.dumps(st.session_state.law_db,ensure_ascii=False,indent=2)
                st.download_button("⬇️ JSON",d.encode("utf-8"),"law_db.json","application/json")
        with col_clr:
            if st.button("🗑️ مسح القاعدة"):
                st.session_state.law_db=[]
                save_law()
                st.success("✅ تم المسح")

# ── TAB 9: الإعدادات ────────────────────────
with t_settings:
    st.subheader("⚙️ الإعدادات")

    # خلفية الموقع
    st.markdown("### 🖼️ خلفية الموقع")
    bg_file=st.file_uploader("ارفع صورة الخلفية (PNG/JPG)",
                              type=["png","jpg","jpeg"],key="bg_up")
    if bg_file:
        raw=_bytes(bg_file)
        ext=(bg_file.name or "").rsplit(".",1)[-1].lower()
        mime=f"image/{'jpeg' if ext in ('jpg','jpeg') else 'png'}"
        b64=base64.b64encode(raw).decode()
        full_b64=f"data:{mime};base64,{b64}"
        st.session_state.bg_b64=b64
        with open(BG_FILE,"w") as f:
            f.write(b64)
        set_bg(b64)
        st.success("✅ تم تطبيق الخلفية")
        st.image(raw,width=200)

    if st.session_state.bg_b64:
        if st.button("🗑️ إزالة الخلفية"):
            st.session_state.bg_b64=""
            if os.path.exists(BG_FILE): os.remove(BG_FILE)
            st.rerun()

    st.markdown("---")

    # تصدير كل شيء
    st.markdown("### 📦 تصدير كامل")
    if st.button("📦 تصدير أرشيف كامل",use_container_width=True):
        export={
            "memory":st.session_state.memory,
            "law_db":st.session_state.law_db,
            "settings":{
                "ai_provider":st.session_state.ai_provider,
                "ai_model":st.session_state.ai_model,
                "ai_format":st.session_state.ai_format,
                "ai_endpoint":st.session_state.ai_endpoint,
            },
            "sessions_count":len(list_sessions()),
            "exported_at":datetime.now().isoformat(),
            "version":"2.0",
        }
        d=json.dumps(export,ensure_ascii=False,indent=2)
        st.download_button("⬇️ تحميل الأرشيف",
            d.encode("utf-8"),
            f"fuehrer_backup_{datetime.now().strftime('%Y%m%d')}.json",
            "application/json")

    st.markdown("---")
    st.markdown("### ℹ️ معلومات النظام")
    st.markdown(f"""
| البيان | القيمة |
|--------|--------|
| النموذج | {st.session_state.ai_provider} |
| الذاكرة | {len(st.session_state.memory)} سجل |
| القانون | {len(st.session_state.law_db)} مادة |
| الجلسات | {len(list_sessions())} جلسة |
| المستندات | {len(st.session_state.docs)} ملف |
""")

st.markdown('<hr><p style="text-align:center;color:#303848;font-size:11px">Führer v2.0 | سري • دائم • أي نموذج ذكاء</p>', unsafe_allow_html=True)
